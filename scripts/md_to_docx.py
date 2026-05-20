"""Convert phase1_采集方案.md to .docx with proper table rendering."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"Q:/Cosmetic/docs/phase1_采集方案.md")
DST = Path(r"Q:/Cosmetic/docs/phase1_采集方案.docx")


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cn_font(run, size=None, bold=None):
    run.font.name = "微软雅黑"
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    rfonts.set(qn("w:ascii"), "微软雅黑")
    rfonts.set(qn("w:hAnsi"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_para(doc, text, size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    return p


def parse_inline(text):
    """Strip markdown emphasis markers; return plain segments with bold flag."""
    parts = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    cleaned = []
    for seg, b in parts:
        seg = seg.replace("`", "")
        cleaned.append((seg, b))
    return cleaned


def add_rich_para(doc, text, size=11, base_bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    for seg, b in parse_inline(text):
        run = p.add_run(seg)
        set_cn_font(run, size=size, bold=base_bold or b)
    return p


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            for seg, b in parse_inline(cell_text):
                run = p.add_run(seg)
                set_cn_font(run, size=10, bold=(i == 0) or b)
            if i == 0:
                set_cell_bg(cell, "4472C4")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table


def parse_table_block(lines, start):
    """Parse a markdown table starting at `start`. Return (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row_line = lines[i].strip()
        cells = [c.strip() for c in row_line.strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c) for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def convert():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    i = 0
    in_code = False
    code_buf = []
    while i < len(md):
        line = md[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                rpr = run._element.get_or_add_rPr()
                rfonts = OxmlElement("w:rFonts")
                rfonts.set(qn("w:ascii"), "Consolas")
                rfonts.set(qn("w:hAnsi"), "Consolas")
                rpr.append(rfonts)
                set_cell_bg_para(p, "F2F2F2")
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("# "):
            add_rich_para(doc, stripped[2:], size=22, base_bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
        elif stripped.startswith("## "):
            add_rich_para(doc, stripped[3:], size=16, base_bold=True)
        elif stripped.startswith("### "):
            add_rich_para(doc, stripped[4:], size=13, base_bold=True)
        elif stripped.startswith("> "):
            add_rich_para(doc, "  " + stripped[2:], size=10, base_bold=False)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("|") and i + 1 < len(md) and "---" in md[i + 1]:
            rows, new_i = parse_table_block(md, i)
            if rows:
                add_table(doc, rows)
            i = new_i
            continue
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            for seg, b in parse_inline(stripped[2:]):
                run = p.add_run(seg)
                set_cn_font(run, size=11, bold=b)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\.\s", "", stripped)
            for seg, b in parse_inline(content):
                run = p.add_run(seg)
                set_cn_font(run, size=11, bold=b)
        elif stripped == "":
            doc.add_paragraph()
        else:
            add_rich_para(doc, stripped, size=11)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


def set_cell_bg_para(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


if __name__ == "__main__":
    convert()
